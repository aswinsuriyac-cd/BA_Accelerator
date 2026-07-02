from io import BytesIO
from typing import Literal

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
import html

from app.schemas.generator_schema import GeneratorOutput, UserStoryRow
import re

ExportFormat = Literal["xlsx", "docx", "pdf"]

def _clean_text(text: str | None) -> str:
    if not text:
        return ""
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', str(text))

EXCEL_HEADERS = [
    "S.no",
    "Epic",
    "Features",
    "US ID",
    "US Summary",
    "User Story Description",
    "Acceptance Criteria",
    "Business Rules",
    "Dependencies",
    "State",
    "Comments",
    "Reference Link",
]


def _block(items: list[str]) -> str:
    return "\n\n".join(item.strip() for item in items if item and item.strip())


def _story_row_values(story: UserStoryRow) -> list[str | int]:
    return [
        _clean_text(story.serial_number),
        _clean_text(story.epic),
        _clean_text(story.feature),
        _clean_text(story.us_id),
        _clean_text(story.us_summary),
        _clean_text(story.user_story_description),
        _clean_text(_block(story.acceptance_criteria)),
        _clean_text(_block(story.business_rules)),
        _clean_text(_block(story.dependencies)),
        _clean_text(story.state),
        _clean_text(story.comments),
        _clean_text(story.reference_link),
    ]


def build_excel_export(output: GeneratorOutput) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "AI-US"

    sheet.append(EXCEL_HEADERS)
    for story in output.stories:
        sheet.append(_story_row_values(story))

    for cell in sheet[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(vertical="top", wrap_text=True)

    column_widths = {
        "A": 8,
        "B": 18,
        "C": 24,
        "D": 14,
        "E": 28,
        "F": 48,
        "G": 42,
        "H": 32,
        "I": 32,
        "J": 12,
        "K": 18,
        "L": 24,
    }
    for column, width in column_widths.items():
        sheet.column_dimensions[column].width = width

    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def build_docx_export(output: GeneratorOutput) -> bytes:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    document.add_heading(_clean_text(output.document_title), level=0)
    document.add_paragraph(f"Story ID Prefix: {_clean_text(output.story_id_prefix)}")
    document.add_paragraph(f"Total Stories: {len(output.stories)}")

    for story in output.stories:
        document.add_heading(f"{_clean_text(story.us_id)} - {_clean_text(story.us_summary)}", level=1)
        document.add_paragraph(f"Epic: {_clean_text(story.epic)}")
        document.add_paragraph(f"Feature: {_clean_text(story.feature)}")
        document.add_paragraph(f"State: {_clean_text(story.state)}")

        story_heading = document.add_paragraph()
        story_heading.add_run("User Story Description: ").bold = True
        story_heading.add_run(_clean_text(story.user_story_description))

        document.add_paragraph("Acceptance Criteria", style="Heading 2")
        for item in story.acceptance_criteria:
            document.add_paragraph(_clean_text(item), style="List Bullet")

        if story.business_rules:
            document.add_paragraph("Business Rules", style="Heading 2")
            for item in story.business_rules:
                document.add_paragraph(_clean_text(item), style="List Bullet")

        if story.dependencies:
            document.add_paragraph("Dependencies", style="Heading 2")
            for item in story.dependencies:
                document.add_paragraph(_clean_text(item), style="List Bullet")

        if story.comments:
            comment_paragraph = document.add_paragraph()
            comment_paragraph.add_run("Comments: ").bold = True
            comment_paragraph.add_run(_clean_text(story.comments))

        if story.reference_link:
            reference_paragraph = document.add_paragraph()
            reference_paragraph.add_run("Reference Link: ").bold = True
            reference_paragraph.add_run(_clean_text(story.reference_link))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf_export(output: GeneratorOutput) -> bytes:
    buffer = BytesIO()
    styles = getSampleStyleSheet()
    document = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=24,
        rightMargin=24,
        topMargin=24,
        bottomMargin=24,
    )

    cell_style = ParagraphStyle('TableCell', parent=styles['Normal'], fontSize=8, leading=10)

    story_table_rows = [["US ID", "Summary", "Epic / Feature", "Story", "Acceptance Criteria"]]
    for story in output.stories:
        story_table_rows.append(
            [
                Paragraph(html.escape(story.us_id), cell_style),
                Paragraph(html.escape(story.us_summary), cell_style),
                Paragraph(f"{html.escape(story.epic)}<br/>{html.escape(story.feature)}", cell_style),
                Paragraph(html.escape(story.user_story_description), cell_style),
                Paragraph(html.escape(_block(story.acceptance_criteria)).replace('\n', '<br/>'), cell_style),
            ]
        )

    elements = [
        Paragraph(output.document_title, styles["Title"]),
        Spacer(1, 8),
        Paragraph(f"Story ID Prefix: {output.story_id_prefix}", styles["Normal"]),
        Spacer(1, 12),
        Table(
            story_table_rows,
            colWidths=[60, 120, 120, 240, 220],
            repeatRows=1,
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E2F3")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F9FC")]),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            ),
        ),
    ]

    document.build(elements)
    return buffer.getvalue()


def build_export_bytes(output: GeneratorOutput, output_format: ExportFormat) -> bytes:
    if output_format == "xlsx":
        return build_excel_export(output)
    if output_format == "docx":
        return build_docx_export(output)
    if output_format == "pdf":
        return build_pdf_export(output)
    raise ValueError(f"Unsupported export format: {output_format}")


def export_media_type(output_format: ExportFormat) -> str:
    if output_format == "xlsx":
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if output_format == "docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return "application/pdf"
