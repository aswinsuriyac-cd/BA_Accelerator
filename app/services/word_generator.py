import os

from docx import Document
from docx.shared import Pt

from app.models.brd_model import BRDModel


class WordGenerator:

    def generate(self, brd: BRDModel):

        document = Document()

        # ------------------------------
        # Title
        # ------------------------------

        heading = document.add_heading(
            "BUSINESS REQUIREMENTS DOCUMENT",
            level=0
        )

        heading.runs[0].font.size = Pt(22)

        document.add_paragraph()

        # ------------------------------
        # Project
        # ------------------------------

        document.add_heading("Project", level=1)

        document.add_paragraph(brd.project_name)

        # ------------------------------

        document.add_heading("Business Objective", level=1)

        document.add_paragraph(brd.business_objective)

        # ------------------------------

        document.add_heading("Scope", level=1)

        document.add_paragraph(brd.scope)

        # ------------------------------

        document.add_heading("Stakeholders", level=1)

        for stakeholder in brd.stakeholders:

            document.add_paragraph(
                stakeholder,
                style="List Bullet"
            )

        # ------------------------------

        document.add_heading(
            "Functional Requirements",
            level=1
        )

        for i, fr in enumerate(
            brd.functional_requirements,
            start=1
        ):

            document.add_paragraph(
                f"FR-{i:03d} : {fr}",
                style="List Bullet"
            )

        # ------------------------------

        document.add_heading(
            "Non Functional Requirements",
            level=1
        )

        for i, nfr in enumerate(
            brd.non_functional_requirements,
            start=1
        ):

            document.add_paragraph(
                f"NFR-{i:03d} : {nfr}",
                style="List Bullet"
            )

        # ------------------------------

        document.add_heading(
            "Constraints",
            level=1
        )

        for item in brd.constraints:

            document.add_paragraph(
                item,
                style="List Bullet"
            )

        # ------------------------------

        document.add_heading(
            "Assumptions",
            level=1
        )

        for item in brd.assumptions:

            document.add_paragraph(
                item,
                style="List Bullet"
            )

        # ------------------------------

        document.add_heading(
            "Dependencies",
            level=1
        )

        for item in brd.dependencies:

            document.add_paragraph(
                item,
                style="List Bullet"
            )

        # ------------------------------

        document.add_heading(
            "Risks",
            level=1
        )

        for item in brd.risks:

            document.add_paragraph(
                item,
                style="List Bullet"
            )

        output_dir = "output"

        os.makedirs(output_dir, exist_ok=True)

        filename = f"{brd.project_name}.docx"

        path = os.path.join(output_dir, filename)

        document.save(path)

        return path