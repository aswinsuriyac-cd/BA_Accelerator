import io
import mimetypes
from csv import reader
from email import policy
from email.parser import BytesParser
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from zipfile import BadZipFile
from openpyxl import load_workbook
from openpyxl.utils.exceptions import InvalidFileException
from google.genai import types

from app.config import settings
from app.services.gemini_service import generate_content_with_fallback

_IMAGE_MIME_TYPES = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "gif": "image/gif",
    "bmp": "image/bmp",
    "tiff": "image/tiff",
    "webp": "image/webp",
}

_OCR_PROMPT = (
    "Extract all visible text from this image, exactly as it appears, "
    "preserving reading order and line breaks. Do not summarize, describe, "
    "or add commentary — output only the raw extracted text. If there is "
    "no legible text, output nothing."
)


def parse_txt(file_bytes: bytes) -> str:
    """Parse TXT or Markdown file content."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def parse_csv(file_bytes: bytes) -> str:
    """Parse CSV file content into pipe-delimited rows."""
    text = parse_txt(file_bytes)
    rows = []
    for row in reader(io.StringIO(text)):
        row_text = " | ".join(cell.strip() for cell in row if cell and cell.strip())
        if row_text:
            rows.append(row_text)
    return "\n".join(rows)


def parse_email(file_bytes: bytes) -> str:
    """Parse EML email content into readable text."""
    message = BytesParser(policy=policy.default).parsebytes(file_bytes)
    parts = [
        f"Subject: {message.get('subject', '')}",
        f"From: {message.get('from', '')}",
        f"To: {message.get('to', '')}",
    ]

    body = message.get_body(preferencelist=("plain", "html"))
    if body is not None:
        parts.append(body.get_content())

    return "\n".join(part for part in parts if part and part.strip())


def parse_pdf(file_bytes: bytes) -> str:
    """Parse PDF file content."""
    pdf_file = io.BytesIO(file_bytes)
    try:
        reader = PdfReader(pdf_file)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                raise ValueError("PDF is password-protected and could not be read.")
        text_content = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_content.append(text)
        return "\n".join(text_content)
    except PdfReadError as e:
        raise ValueError(f"Could not parse PDF: {e}")


def parse_docx(file_bytes: bytes) -> str:
    """Parse DOCX file content."""
    docx_file = io.BytesIO(file_bytes)
    try:
        doc = Document(docx_file)
    except (PackageNotFoundError, BadZipFile) as e:
        raise ValueError(f"Could not parse DOCX: {e}")

    text_content = []
    for para in doc.paragraphs:
        if para.text:
            text_content.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text]
            if row_text:
                text_content.append(" | ".join(row_text))

    return "\n".join(text_content)


def parse_excel(file_bytes: bytes) -> str:
    """Parse XLSX/XLSM file content, sheet by sheet."""
    excel_file = io.BytesIO(file_bytes)
    try:
        # read_only + data_only: streams rows efficiently and returns
        # computed values instead of formula strings (e.g. "150" not "=A1+A2")
        workbook = load_workbook(excel_file, read_only=True, data_only=True)
    except (InvalidFileException, BadZipFile) as e:
        raise ValueError(f"Could not parse Excel file: {e}")

    text_content = []
    try:
        for sheet in workbook.worksheets:
            sheet_rows = []
            for row in sheet.iter_rows(values_only=True):
                # Skip fully empty rows
                if all(cell is None for cell in row):
                    continue
                row_text = " | ".join(
                    str(cell).strip() for cell in row if cell is not None
                )
                if row_text:
                    sheet_rows.append(row_text)

            if sheet_rows:
                text_content.append(f"## Sheet: {sheet.title}")
                text_content.extend(sheet_rows)
    finally:
        workbook.close()

    return "\n".join(text_content)


def parse_image(file_bytes: bytes, filename: str | None = None) -> str:
    """Parse image file content via Gemini vision (no OS OCR dependency)."""
    mime_type = None
    if filename:
        mime_type, _ = mimetypes.guess_type(filename)

    if mime_type not in _IMAGE_MIME_TYPES.values():
        # fall back to sniffing magic bytes if extension-based guess failed
        if file_bytes.startswith(b"\x89PNG\r\n\x1a\n"):
            mime_type = "image/png"
        elif file_bytes.startswith(b"\xff\xd8\xff"):
            mime_type = "image/jpeg"
        elif file_bytes.startswith((b"GIF87a", b"GIF89a")):
            mime_type = "image/gif"
        elif file_bytes.startswith(b"RIFF") and file_bytes[8:12] == b"WEBP":
            mime_type = "image/webp"
        elif file_bytes.startswith(b"BM"):
            mime_type = "image/bmp"
        elif file_bytes.startswith((b"II*\x00", b"MM\x00*")):
            mime_type = "image/tiff"
        else:
            raise ValueError("Could not determine image format.")

    try:
        response = generate_content_with_fallback(
            model=settings.model_name,
            contents=[
                types.Part.from_bytes(data=file_bytes, mime_type=mime_type),
                _OCR_PROMPT,
            ],
            config=types.GenerateContentConfig(),
        )
    except Exception as e:
        raise RuntimeError(f"Image text extraction failed: {e}")

    return (response.text or "").strip()


def parse_document(filename: str, content: bytes) -> str:
    """Unified entry point to parse documents based on extension."""
    if "." not in filename:
        raise ValueError(f"Filename '{filename}' has no extension.")

    ext = filename.lower().rsplit(".", 1)[-1]
    if ext in ("txt", "md", "markdown"):
        return parse_txt(content)
    elif ext == "csv":
        return parse_csv(content)
    elif ext == "eml":
        return parse_email(content)
    elif ext == "pdf":
        return parse_pdf(content)
    elif ext == "docx":
        return parse_docx(content)
    elif ext in ("xlsx", "xlsm"):
        return parse_excel(content)
    elif ext in _IMAGE_MIME_TYPES:
        return parse_image(content, filename)
    else:
        raise ValueError(
            f"Unsupported file format: .{ext}. Supported formats are: "
            f".txt, .md, .markdown, .csv, .eml, .pdf, .docx, .xlsx, .xlsm, "
            f"{', '.join('.' + e for e in _IMAGE_MIME_TYPES)}"
        )
